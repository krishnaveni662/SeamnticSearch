"""
Single-file semantic-search app using:
 - local CPU extraction (pdfplumber, pytesseract, python-docx)
 - Hugging Face Inference API for text-generation (Qwen endpoint) and embeddings
 - Milvus for vector storage and semantic search
 - simple RBAC per-document
 - Streamlit UI

Requirements (install with pip):
  pip install streamlit pymilvus requests pdfplumber pytesseract python-docx pillow sentence-transformers python-dotenv

Notes:
 - Set HUGGINGFACE_API_TOKEN in your environment or a .env file.
 - Adjust HF_TEXT_MODEL and HF_EMBEDDING_MODEL env vars if needed.
 - Tesseract must be installed on Windows and in PATH for pytesseract OCR.
"""
import os
import io
import json
import tempfile
import base64
from pathlib import Path
from typing import Tuple, List, Optional
import logging
import threading

import streamlit as st
import requests
from PIL import Image

# Document extraction libs
import pdfplumber
import pytesseract
from docx import Document

# Milvus
from pymilvus import connections, Collection, FieldSchema, CollectionSchema, DataType, utility

# Sentence-transformers (fallback/local embedding if desired)
from sentence_transformers import SentenceTransformer

# Load env
from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("semantic_search_app")

# ---------- Configuration ----------
HF_API_TOKEN = os.getenv("HUGGINGFACE_API_TOKEN", "")
HF_TEXT_MODEL = os.getenv("HF_TEXT_MODEL", "Qwen/Qwen2.5-VL-7B-Instruct")
HF_EMBEDDING_MODEL = os.getenv("HF_EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
MILVUS_HOST = os.getenv("MILVUS_HOST", "localhost")
MILVUS_PORT = int(os.getenv("MILVUS_PORT", 19530))
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "semantic_search")
EMBEDDING_DIM_OVERRIDE = os.getenv("EMBEDDING_DIM")  # optional
DEVICE = os.getenv("DEVICE", "cpu")
MAX_WORKERS = int(os.getenv("MAX_WORKERS", "2"))

HEADERS = {"Authorization": f"Bearer {HF_API_TOKEN}"} if HF_API_TOKEN else {}

# ---------- Utilities: Hugging Face Inference calls ----------
def hf_text_generate(prompt: str, model: str = HF_TEXT_MODEL, timeout: int = 60) -> str:
    """
    Call HF text-generation endpoint. Returns string output (or fallback to prompt if no token).
    """
    if not HF_API_TOKEN:
        logger.warning("No HF token; returning original text (no remote generation).")
        return prompt[:4000]

    url = f"https://api-inference.huggingface.co/models/{model}"
    payload = {"inputs": prompt, "options": {"wait_for_model": True, "use_cache": False}}
    try:
        resp = requests.post(url, headers=HEADERS, json=payload, timeout=timeout)
        resp.raise_for_status()
        out = resp.json()
        # HF generative endpoints usually return list of dicts with 'generated_text' or a single string.
        if isinstance(out, list) and len(out) > 0:
            text = out[0].get("generated_text") or out[0].get("text") or str(out[0])
            return text
        if isinstance(out, dict) and "generated_text" in out:
            return out["generated_text"]
        return str(out)
    except Exception as e:
        logger.exception("HF text generation failed: %s", e)
        return prompt[:4000]

def hf_get_embedding(text: str, model: str = HF_EMBEDDING_MODEL, timeout: int = 30) -> Optional[List[float]]:
    """
    Call HF inference endpoint for embeddings. Expects model to return list of floats.
    Falls back to local sentence-transformers if no token or API fails.
    """
    if HF_API_TOKEN:
        url = f"https://api-inference.huggingface.co/models/{model}"
        payload = {"inputs": text, "options": {"wait_for_model": True}}
        try:
            resp = requests.post(url, headers=HEADERS, json=payload, timeout=timeout)
            resp.raise_for_status()
            out = resp.json()
            # API often returns a list of floats or nested; try to locate first list of floats
            if isinstance(out, list) and all(isinstance(x, (float, int)) for x in out):
                return [float(x) for x in out]
            if isinstance(out, list) and len(out) > 0 and isinstance(out[0], list):
                return [float(x) for x in out[0]]
            # Some HF embedding endpoints return {"embedding": [...]} or similar
            if isinstance(out, dict) and "embedding" in out:
                return [float(x) for x in out["embedding"]]
            logger.warning("Unexpected embedding response shape: %s", out)
        except Exception as e:
            logger.exception("HF embedding call failed: %s", e)

    # Fallback local embedding (CPU friendly)
    try:
        logger.info("Using local sentence-transformers fallback for embeddings.")
        local_model = SentenceTransformer(HF_EMBEDDING_MODEL, device=DEVICE)
        vec = local_model.encode(text, convert_to_numpy=True).tolist()
        return [float(x) for x in vec]
    except Exception as e:
        logger.exception("Local embedding fallback failed: %s", e)
        return None

# ---------- Document extraction helpers ----------
def extract_text_from_image(img: Image.Image) -> str:
    try:
        text = pytesseract.image_to_string(img)
        return text or ""
    except Exception as e:
        logger.exception("Image OCR failed: %s", e)
        return ""

def extract_text_from_pdf(path_or_bytes) -> str:
    texts = []
    try:
        if isinstance(path_or_bytes, (bytes, io.BytesIO)):
            fp = io.BytesIO(path_or_bytes if isinstance(path_or_bytes, bytes) else path_or_bytes.getvalue())
            pdf = pdfplumber.open(fp)
        else:
            pdf = pdfplumber.open(path_or_bytes)
        with pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if not page_text.strip():
                    try:
                        pil = page.to_image(resolution=150).original
                        page_text = pytesseract.image_to_string(pil)
                    except Exception:
                        page_text = ""
                texts.append(page_text)
    except Exception as e:
        logger.exception("PDF extraction error: %s", e)
    return "\n---PAGE---\n".join(texts)

def extract_text_from_docx(path_or_bytes) -> str:
    try:
        if isinstance(path_or_bytes, (bytes, io.BytesIO)):
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp.write(path_or_bytes if isinstance(path_or_bytes, bytes) else path_or_bytes.getvalue())
            tmp.close()
            doc = Document(tmp.name)
        else:
            doc = Document(path_or_bytes)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.exception("DOCX extraction error: %s", e)
        return ""

def extract_text_from_txt(path_or_bytes) -> str:
    try:
        if isinstance(path_or_bytes, (bytes, io.BytesIO)):
            data = path_or_bytes.read() if hasattr(path_or_bytes, "read") else path_or_bytes
            return data.decode("utf-8", errors="ignore")
        with open(path_or_bytes, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.exception("Text extraction error: %s", e)
        return ""

# ---------- Milvus client wrapper ----------
class MilvusClient:
    def __init__(self, host=MILVUS_HOST, port=MILVUS_PORT, collection_name=COLLECTION_NAME):
        connections.connect(host=host, port=port)
        self.collection_name = collection_name
        self.collection = None
        # create lazily once we know dimension
    def create_collection(self, dim: int):
        if utility.has_collection(self.collection_name):
            self.collection = Collection(self.collection_name)
            return self.collection
        id_field = FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True)
        emb_field = FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=dim)
        content_field = FieldSchema(name="content", dtype=DataType.VARCHAR, max_length=65535)
        doc_type_field = FieldSchema(name="doc_type", dtype=DataType.VARCHAR, max_length=64)
        role_field = FieldSchema(name="role", dtype=DataType.VARCHAR, max_length=64)
        schema = CollectionSchema(fields=[id_field, emb_field, content_field, doc_type_field, role_field],
                                  description="Semantic search collection")
        self.collection = Collection(name=self.collection_name, schema=schema)
        index_params = {"index_type": "IVF_FLAT", "metric_type": "COSINE", "params": {"nlist": 128}}
        self.collection.create_index(field_name="embedding", index_params=index_params)
        self.collection.load()
        return self.collection
    def insert(self, embeddings: List[List[float]], contents: List[str], doc_types: List[str], roles: List[str]):
        if not embeddings:
            return None
        dim = len(embeddings[0])
        if not self.collection:
            self.create_collection(dim)
        entities = [embeddings, contents, doc_types, roles]
        mr = self.collection.insert(entities)
        self.collection.flush()
        return mr
    def search(self, query_embedding: List[float], role: str = None, top_k: int = 5):
        if not self.collection:
            logger.warning("Collection not found or empty.")
            return []
        expr = f'role == "{role}"' if role else None
        search_params = {"metric_type": "COSINE", "params": {"nprobe": 10}}
        results = self.collection.search([query_embedding], "embedding", param=search_params, limit=top_k, expr=expr)
        # results is list of hits per query
        return results

# ---------- RBAC ----------
class RBAC:
    def __init__(self):
        self.roles = {
            "admin": ["read", "write", "delete"],
            "user": ["read", "write"],
            "viewer": ["read"]
        }
    def has_permission(self, role: str, action: str) -> bool:
        return action in self.roles.get(role, [])

# ---------- Batch processing ----------
class Processor:
    def __init__(self, hf_text_model=HF_TEXT_MODEL, hf_embedding_model=HF_EMBEDDING_MODEL):
        self.hf_text_model = hf_text_model
        self.hf_embedding_model = hf_embedding_model
    def process_document(self, content_obj, file_type: str) -> Tuple[str, Optional[List[float]]]:
        """
        content_obj: path str, bytes, BytesIO, or PIL.Image
        file_type: 'pdf','docx','txt','image'
        Returns processed_text (structured via HF) and embedding vector.
        """
        raw_text = ""
        # extract raw text
        if file_type == "image":
            if isinstance(content_obj, Image.Image):
                raw_text = extract_text_from_image(content_obj)
            elif isinstance(content_obj, (bytes, io.BytesIO)):
                try:
                    img = Image.open(io.BytesIO(content_obj if isinstance(content_obj, bytes) else content_obj.getvalue()))
                    raw_text = extract_text_from_image(img)
                except Exception:
                    raw_text = ""
            else:
                # assume path
                img = Image.open(content_obj)
                raw_text = extract_text_from_image(img)
        elif file_type == "pdf":
            raw_text = extract_text_from_pdf(content_obj)
        elif file_type == "docx":
            raw_text = extract_text_from_docx(content_obj)
        elif file_type == "text":
            raw_text = extract_text_from_txt(content_obj)
        else:
            raw_text = str(content_obj)

        # Create a prompt for HF generative model to structure/extract tables/charts/key info
        prompt = (
            "Extract and structure the content of the following document. "
            "Provide:\n1) Clean text summary\n2) List of detected tables with rows (if any)\n"
            "3) Detected charts with type and key values (if any)\n4) Key metadata/key-value pairs.\n\n"
            "Document content:\n\n" + (raw_text[:30000])  # truncate large docs for safety
        )

        structured = hf_text_generate(prompt, model=self.hf_text_model)
        embedding = hf_get_embedding(structured)
        return structured, embedding

# ---------- Streamlit UI (single-file) ----------
st.set_page_config(page_title="Semantic Search (HF + Milvus)", layout="wide")
st.title("Semantic Search — Hugging Face Inference + Milvus (single-file)")

# Initialize components
rbac = RBAC()
processor = Processor()
milvus = MilvusClient()
# ensure tmp dir
os.makedirs(r"c:\milvus\tmp_upload", exist_ok=True)

# Sidebar: user + settings
with st.sidebar:
    st.header("Settings & User")
    username = st.text_input("Username", value="alice")
    role = st.selectbox("Role", ["admin", "user", "viewer"], index=1)
    st.write("HF token provided:" , bool(HF_API_TOKEN))
    st.write("HF text model:", HF_TEXT_MODEL)
    st.write("HF embedding model:", HF_EMBEDDING_MODEL)
    st.write("Milvus:", f"{MILVUS_HOST}:{MILVUS_PORT}")
    st.write("CPU-friendly workers:", MAX_WORKERS)
    st.markdown("---")
    st.markdown("Permissions:")
    st.json(rbac.roles)

tabs = st.tabs(["Upload & Index", "Batch Index Directory", "Search"])

# Upload & Index
with tabs[0]:
    st.header("Upload and index a document")
    upload = st.file_uploader("Upload file (pdf, docx, txt, image)", type=["pdf","docx","txt","png","jpg","jpeg","tiff","bmp"])
    assign_role = st.selectbox("Assign role to this document", ["user","viewer","admin"], index=0)
    if upload:
        if not rbac.has_permission(role, "write"):
            st.error("You don't have write permission.")
        else:
            st.info("Saving and processing...")
            tmp_path = Path(r"c:\milvus\tmp_upload") / upload.name
            with open(tmp_path, "wb") as f:
                f.write(upload.getbuffer())
            ext = tmp_path.suffix.lower()
            file_type = "text"
            if ext == ".pdf":
                file_type = "pdf"
                content_obj = str(tmp_path)
            elif ext == ".docx":
                file_type = "docx"
                content_obj = str(tmp_path)
            elif ext in [".txt", ".md"]:
                file_type = "text"
                content_obj = str(tmp_path)
            else:
                file_type = "image"
                content_obj = Image.open(tmp_path)
            structured, emb = processor.process_document(content_obj, file_type)
            if emb is None:
                st.error("Failed to generate embedding.")
            else:
                milvus.insert([emb], [structured], [file_type], [assign_role])
                st.success(f"Indexed {upload.name} as role={assign_role}")
                st.text_area("Structured extraction (preview)", structured, height=300)

# Batch Index Directory
with tabs[1]:
    st.header("Batch index a directory")
    dir_path = st.text_input("Directory path to index", value=r"c:\milvus\data")
    batch_role = st.selectbox("Role for indexed docs", ["user","viewer","admin"], index=0)
    if st.button("Run batch index"):
        if not rbac.has_permission(role, "write"):
            st.error("No write permission.")
        else:
            p = Path(dir_path)
            if not p.exists() or not p.is_dir():
                st.error("Directory not found.")
            else:
                files = [f for f in p.rglob("*") if f.suffix.lower() in {".pdf",".docx",".txt",".md",".png",".jpg",".jpeg",".tiff",".bmp"}]
                st.info(f"Found {len(files)} files.")
                progress = st.progress(0)
                results = []
                total = len(files)
                for i, f in enumerate(files, start=1):
                    ext = f.suffix.lower()
                    if ext == ".pdf":
                        ft = "pdf"
                        content = str(f)
                    elif ext == ".docx":
                        ft = "docx"
                        content = str(f)
                    elif ext in [".txt",".md"]:
                        ft = "text"
                        content = str(f)
                    else:
                        ft = "image"
                        content = Image.open(f)
                    structured, emb = processor.process_document(content, ft)
                    if emb:
                        milvus.insert([emb], [structured], [ft], [batch_role])
                        results.append({"file": str(f), "status": "indexed"})
                    else:
                        results.append({"file": str(f), "status": "failed"})
                    progress.progress(i/total)
                st.write(results)
                st.success("Batch indexing completed.")

# Search
with tabs[2]:
    st.header("Semantic Search")
    query = st.text_area("Enter search query (text)", height=120)
    top_k = st.slider("Top K", min_value=1, max_value=20, value=5)
    role_filter = st.selectbox("Filter results by document role", ["", "user", "viewer", "admin"], index=0)
    if st.button("Search"):
        if not rbac.has_permission(role, "read"):
            st.error("No read permission.")
        else:
            if not query.strip():
                st.warning("Enter query text.")
            else:
                st.info("Generating embedding for query...")
                q_emb = hf_get_embedding(query)
                if not q_emb:
                    st.error("Failed to get query embedding.")
                else:
                    res = milvus.search(q_emb, role=role_filter if role_filter else None, top_k=top_k)
                    hits = []
                    # pymilvus search returns list of topk result sets (one per input)
                    if not res:
                        st.info("No results (collection empty or role filter).")
                    else:
                        for hit in res[0]:
                            # each hit has .id, .distance, .entity
                            content_preview = ""
                            doc_type = ""
                            try:
                                ent = hit.entity
                                content_preview = ent.get("content", "")[:1200] if ent and "content" in ent else ""
                                doc_type = ent.get("doc_type", "")
                            except Exception:
                                pass
                            hits.append({"id": getattr(hit, "id", None), "score": getattr(hit, "distance", None) or getattr(hit, "score", None),
                                         "content_preview": content_preview, "doc_type": doc_type})
                        if not hits:
                            st.info("No matching documents.")
                        else:
                            for i, h in enumerate(hits, start=1):
                                st.subheader(f"{i}. score={h['score']}")
                                st.code(h["content_preview"])
                                st.caption(f"doc_type: {h['doc_type']}  id: {h['id']}")

# Footer note
st.markdown("---")
st.caption("Notes: This app uses local OCR/text extraction and calls Hugging Face Inference API for generation/embeddings when HUGGINGFACE_API_TOKEN is set. For large-scale/production use, adapt batching, error handling, and credentials management.")